import nbformat
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


with open("NLTK_EXPLORE.ipynb", "r", encoding="utf-8") as f:
    nb = nbformat.read(f, as_version=4)

doc = Document()
doc.add_heading('Notebook Export', level=1)

previous_cell_was_code_no_output = False

for cell in nb.cells:
    if cell.cell_type == 'markdown':

        p = doc.add_paragraph()
        run = p.add_run(cell.source)
        run.bold = True
        run.font.size = Pt(12)
        previous_cell_was_code_no_output = False

    elif cell.cell_type == 'code':
        has_output = bool(cell.outputs)


        if has_output:
            doc.add_paragraph("CODE:")


        code_p = doc.add_paragraph()
        code_run = code_p.add_run(cell.source)
        code_run.font.name = 'Courier New'
        code_run.font.size = Pt(10)

        if has_output:
            for output in cell.outputs:
                if 'text' in output:
                    doc.add_paragraph("OUTPUT:")
                    out_p = doc.add_paragraph()
                    out_run = out_p.add_run(output['text'])
                    out_run.font.name = 'Courier New'
                    out_run.font.size = Pt(10)

            previous_cell_was_code_no_output = False
        else:

            doc.add_paragraph()
            previous_cell_was_code_no_output = True

doc.save("Notebook_Export.docx")
print("✅ Word document created with code and output.")
